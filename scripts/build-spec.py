#!/usr/bin/env python3
"""
Generate docs/master-spec.docx from docs/master-spec.md.

The markdown file is the source of truth — it lives in the repo, diffs in review, and links
to each module's own spec. The .docx is a shareable export, regenerated on demand.

Usage:
    pip install python-docx
    python scripts/build-spec.py [input.md] [output.docx]

Supports the markdown subset used by the spec: headings, paragraphs, bullet lists, tables,
fenced code blocks, and inline **bold**, `code`, and [links](target).
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

MONO = "Consolas"
LINK_BLUE = "0563C1"

INLINE = re.compile(r"(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|`[^`]+`)")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_hyperlink(paragraph, text, url, mono=False):
    """python-docx has no hyperlink API; build the relationship and run by hand."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(color)
    props.append(underline)
    if mono:
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), MONO)
        fonts.set(qn("w:hAnsi"), MONO)
        props.append(fonts)
    run.append(props)

    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def add_runs(paragraph, text):
    """Render inline markdown into runs on an existing paragraph."""
    for token in INLINE.split(text):
        if not token:
            continue

        match = LINK.fullmatch(token)
        if match:
            label, target = match.group(1), match.group(2)
            mono = label.startswith("`") and label.endswith("`")
            add_hyperlink(paragraph, label.strip("`"), target, mono=mono)
            continue

        if token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
            continue

        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = MONO
            run.font.size = Pt(9.5)
            continue

        paragraph.add_run(token)


def split_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def build(md_path: Path, out_path: Path):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10.5)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank lines and horizontal rules
        if not stripped or stripped == "---":
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            block = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            for code_line in block:
                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(0)
                run = para.add_run(code_line or " ")
                run.font.name = MONO
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            doc.add_paragraph()
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) and is_separator(lines[i + 1]):
            header = split_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1

            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            for idx, text in enumerate(header):
                cell = table.rows[0].cells[idx]
                cell.text = ""
                add_runs(cell.paragraphs[0], text)
                for run in cell.paragraphs[0].runs:
                    run.bold = True

            for row in rows:
                cells = table.add_row().cells
                for idx, text in enumerate(row[: len(header)]):
                    cells[idx].text = ""
                    add_runs(cells[idx].paragraphs[0], text)
            doc.add_paragraph()
            continue

        # Headings
        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            para = doc.add_heading("", 0 if level == 1 else level - 1)
            add_runs(para, heading.group(2))
            i += 1
            continue

        # Bullet list
        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            add_runs(para, bullet.group(1))
            i += 1
            continue

        # Numbered list
        number = re.match(r"^\d+\.\s+(.*)$", stripped)
        if number:
            para = doc.add_paragraph(style="List Number")
            add_runs(para, number.group(1))
            i += 1
            continue

        # Paragraph — join wrapped lines
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt == "---" or nxt.startswith(("#", "|", "```", "- ", "* "))
                    or re.match(r"^\d+\.\s", nxt)):
                break
            buf.append(nxt)
            i += 1
        add_runs(doc.add_paragraph(), " ".join(buf))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parent.parent
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else root / "docs" / "master-spec.md"
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else root / "docs" / "master-spec.docx"
    if not src.exists():
        sys.exit(f"Not found: {src}")
    build(src, dst)
